"""
SEO Content Optimizer
----------------------
Streamlit tool that:
1. Takes a "fresh" (non-optimized) Word doc.
2. Lets the user enter Meta Title / Meta Description / URL, pick a Schema type
   (Article or Product), and fill in a couple of publisher/date fields.
3. On clicking "Optimise":
   - Inserts a Meta Title / Meta Description / URL table at the very top.
   - Wraps every heading's text in the matching HTML tag (<h1>, <h2>, ...).
   - Appends a "Schema Recommendations" section with the chosen schema
     (Article or Product) plus an auto-generated Breadcrumb schema, both as
     JSON-LD <script> blocks — matching the house template.
4. Shows a preview of everything it changed, then lets the user download the
   optimized .docx.
5. Anchor Text & URL Checker: pulls every hyperlink out of the uploaded doc
   and checks anchor text quality + whether the URL is actually live.

An "optimized" reference doc can also be uploaded — it's shown side-by-side
purely for visual comparison, it isn't parsed.

Run with:  streamlit run seo_doc_optimizer.py
Requires:  pip install streamlit python-docx requests
"""

import io
import json
import concurrent.futures
from datetime import datetime
from urllib.parse import urlparse

import requests
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DEFAULT_PUBLISHER_NAME = "Organization"
DEFAULT_PUBLISHER_LOGO = "https://www.mars.com/themes/custom/mars_acss/assets/images/logo-main.svg"
DEFAULT_DATE_PUBLISHED = "To be updated by the team"
GENERIC_ANCHOR_PHRASES = {
    "click here", "here", "read more", "learn more", "this",
    "this link", "link", "more info", "this page", "website", "",
}


# --------------------------------------------------------------------------
# Heading tagging
# --------------------------------------------------------------------------

def extract_headings(doc: Document):
    """Return list of (paragraph_index, level:int, text:str) for every
    paragraph using a built-in Heading style (Heading 1..6) or Title."""
    headings = []
    for i, para in enumerate(doc.paragraphs):
        style_name = (para.style.name or "").strip()
        if style_name == "Title":
            headings.append((i, 1, para.text))
        elif style_name.startswith("Heading"):
            digits = "".join(ch for ch in style_name if ch.isdigit())
            level = int(digits) if digits else 1
            level = min(max(level, 1), 6)
            headings.append((i, level, para.text))
    return headings


def wrap_heading_text(para, level: int):
    """Replace a heading paragraph's runs with a single run whose text is
    wrapped in <hN>...</hN>, keeping the paragraph's heading style."""
    original_text = para.text
    if not original_text.strip():
        return
    bold = para.runs[0].bold if para.runs else None
    italic = para.runs[0].italic if para.runs else None

    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    new_run = para.add_run(f"<h{level}>{original_text}</h{level}>")
    if bold is not None:
        new_run.bold = bold
    if italic is not None:
        new_run.italic = italic


# --------------------------------------------------------------------------
# Meta table (matches the house template: 2-col Table Grid, no extra label)
# --------------------------------------------------------------------------

def set_table_grid_borders(table):
    """Apply visible single-line borders to every cell. Used as a fallback
    when the document has no 'Table Grid' style defined (e.g. docs exported
    from Google Docs), so the meta table still renders with a grid."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "4")
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), "auto")
        borders.append(edge_el)
    tbl_pr.append(borders)


def insert_meta_table_at_top(doc: Document, title: str, description: str, url: str):
    table = doc.add_table(rows=3, cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        set_table_grid_borders(table)
    table.columns[0].width = Inches(1.3)
    table.columns[1].width = Inches(4.5)

    rows_data = [
        ("Meta Title", title),
        ("Meta Description", description),
        ("URL", url),
    ]
    for row, (label, value) in zip(table.rows, rows_data):
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
        for cell in row.cells:
            cell.width = Inches(1.3) if cell is row.cells[0] else Inches(4.5)

    body = doc.element.body
    table_elm = table._tbl
    body.remove(table_elm)
    body.insert(0, table_elm)

    spacer = doc.add_paragraph()
    spacer_elm = spacer._p
    body.remove(spacer_elm)
    body.insert(1, spacer_elm)


# --------------------------------------------------------------------------
# Schema builders
# --------------------------------------------------------------------------

def build_article_schema(title, description, url, publisher_name, publisher_logo, date_published, image_url=""):
    return {
        "@context": "https://schema.org",
        "@type": "Article",
        "mainEntityOfPage": {"@type": "WebPage", "@id": url},
        "headline": title,
        "description": description,
        "image": image_url,
        "publisher": {
            "@type": "Organization",
            "name": publisher_name or DEFAULT_PUBLISHER_NAME,
            "logo": {"@type": "ImageObject", "url": publisher_logo or DEFAULT_PUBLISHER_LOGO},
        },
        "datePublished": date_published or DEFAULT_DATE_PUBLISHED,
    }


def build_product_schema(title, description, url, image_url="", brand="", price="", currency="", availability=""):
    data = {
        "@context": "https://schema.org",
        "@type": "Product",
        "name": title,
        "description": description,
        "image": image_url,
        "url": url,
    }
    if brand:
        data["brand"] = {"@type": "Brand", "name": brand}
    if price or currency or availability:
        offers = {"@type": "Offer", "url": url}
        if price:
            offers["price"] = price
        if currency:
            offers["priceCurrency"] = currency
        if availability:
            offers["availability"] = f"https://schema.org/{availability}"
        data["offers"] = offers
    return data


def build_breadcrumb_schema(url: str):
    parsed = urlparse(url)
    segments = [s for s in parsed.path.split("/") if s]
    items = [{
        "@type": "ListItem",
        "position": 1,
        "name": "Home page",
        "item": f"{parsed.scheme}://{parsed.netloc}/",
    }]
    for i, seg in enumerate(segments):
        is_last = i == len(segments) - 1
        partial_path = "/".join(segments[: i + 1])
        item_url = f"{parsed.scheme}://{parsed.netloc}/{partial_path}"
        name = seg if is_last else seg.replace("-", " ").replace("_", " ").title()
        items.append({
            "@type": "ListItem",
            "position": i + 2,
            "name": name,
            "item": item_url,
        })
    return {"@context": "https://schema.org/", "@type": "BreadcrumbList", "itemListElement": items}


def _add_code_paragraph(doc, text, highlight=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    if highlight:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def _add_json_ld_block(doc, schema_dict, highlight_value=None):
    _add_code_paragraph(doc, '<script type="application/ld+json">')
    json_text = json.dumps(schema_dict, indent=2)
    for line in json_text.split("\n"):
        do_highlight = bool(highlight_value) and highlight_value in line
        _add_code_paragraph(doc, line, highlight=do_highlight)
    _add_code_paragraph(doc, "</script>")


def append_schema_section(doc: Document, schema_type: str, primary_schema: dict,
                           breadcrumb_schema: dict, date_published_value: str):
    doc.add_page_break()

    section_heading = doc.add_paragraph()
    section_heading.add_run("Schema Recommendations").bold = True

    sub1 = doc.add_paragraph()
    sub1.add_run(f"{schema_type} Schema").bold = True
    _add_json_ld_block(doc, primary_schema, highlight_value=date_published_value)

    sub2 = doc.add_paragraph()
    sub2.add_run("Breadcrumb Schema").bold = True
    _add_json_ld_block(doc, breadcrumb_schema)


# --------------------------------------------------------------------------
# Preview helper for the optional reference doc
# --------------------------------------------------------------------------

def read_docx_text_preview(file_bytes: bytes, max_paragraphs: int = 40):
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs[:max_paragraphs]:
        if para.text.strip():
            lines.append(f"[{para.style.name}] {para.text}")
    return "\n".join(lines) if lines else "(no readable text found)"


# --------------------------------------------------------------------------
# Anchor Text & URL Checker
# --------------------------------------------------------------------------

def extract_hyperlinks(doc: Document):
    links = []
    for para in doc.paragraphs:
        try:
            hyperlinks = para.hyperlinks
        except AttributeError:
            hyperlinks = []
        for h in hyperlinks:
            if h.address:
                links.append({"anchor_text": h.text, "url": h.address})
    return links


def classify_anchor(anchor_text: str, url: str):
    text = (anchor_text or "").strip()
    notes = []
    if not text:
        notes.append("Empty/whitespace anchor text")
    else:
        if text.lower() in GENERIC_ANCHOR_PHRASES:
            notes.append("Generic anchor text")
        try:
            netloc = urlparse(url).netloc.lower().replace("www.", "")
            if text.lower().replace("www.", "") == netloc:
                notes.append("Anchor text is bare domain")
        except Exception:
            pass
    return "; ".join(notes) if notes else "OK"


def check_url_status(url: str, timeout: int = 8):
    headers = {"User-Agent": "Mozilla/5.0 (compatible; SEO-Doc-Optimizer-LinkChecker/1.0)"}
    try:
        resp = requests.head(url, headers=headers, timeout=timeout, allow_redirects=True)
        if resp.status_code >= 400:
            resp = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True, stream=True)
            resp.close()
        return resp.status_code
    except requests.exceptions.RequestException:
        return None


def status_label(code):
    if code is None:
        return "Unreachable"
    if 200 <= code < 300:
        return f"OK ({code})"
    if 300 <= code < 400:
        return f"Redirect ({code})"
    return f"Broken ({code})"


# --------------------------------------------------------------------------
# Streamlit UI
# --------------------------------------------------------------------------

st.set_page_config(page_title="SEO Content Optimizer", layout="wide")
st.title("SEO Content Optimizer")
st.caption(
    "Upload a non-optimized doc, fill in the meta details, and generate an "
    "optimized version with a meta table, tagged headings, and schema markup."
)

col_up1, col_up2 = st.columns(2)
with col_up1:
    fresh_file = st.file_uploader(
        "Non-optimized (fresh) document — required", type=["docx"], key="fresh_upload",
    )
with col_up2:
    optimized_file = st.file_uploader(
        "Optimized reference document — optional, for comparison only",
        type=["docx"], key="optimized_upload",
    )

if optimized_file is not None:
    with st.expander("Preview: optimized reference doc (for comparison)"):
        st.text(read_docx_text_preview(optimized_file.getvalue()))

if fresh_file is None:
    st.info("Upload the non-optimized document to continue.")
    st.stop()

fresh_bytes = fresh_file.getvalue()
source_doc_preview = Document(io.BytesIO(fresh_bytes))
found_headings = extract_headings(source_doc_preview)

with st.expander(f"Preview: fresh doc — {len(found_headings)} heading(s) detected"):
    if found_headings:
        for _, level, text in found_headings:
            st.write(f"H{level}: {text}")
    else:
        st.write("No Heading-styled paragraphs found in this document.")

st.divider()
st.subheader("Meta details")

meta_col1, meta_col2 = st.columns(2)
with meta_col1:
    title_input = st.text_input("Title", placeholder="e.g. Best Cat Food for Sensitive Stomachs | NUTRO")
    url_input = st.text_input("URL", placeholder="https://www.example.com/page")
with meta_col2:
    description_input = st.text_area("Description", placeholder="Meta description text...", height=100)

schema_type = st.selectbox("Schema", ["Article", "Product"])

image_input = st.text_input("Image URL (optional, used in schema)", key="image_url")

st.caption("Schema publisher / date fields (used in the JSON-LD)")
sp1, sp2, sp3 = st.columns(3)
with sp1:
    publisher_name_input = st.text_input("Publisher name", value=DEFAULT_PUBLISHER_NAME)
with sp2:
    publisher_logo_input = st.text_input("Publisher logo URL", value=DEFAULT_PUBLISHER_LOGO)
with sp3:
    date_published_input = st.text_input("Date published", value=DEFAULT_DATE_PUBLISHED)

brand_input = price_input = currency_input = availability_input = ""
if schema_type == "Product":
    st.caption("Optional Product schema fields")
    p1, p2, p3, p4 = st.columns(4)
    with p1:
        brand_input = st.text_input("Brand", key="brand")
    with p2:
        price_input = st.text_input("Price", key="price")
    with p3:
        currency_input = st.text_input("Currency (e.g. AUD)", key="currency")
    with p4:
        availability_input = st.selectbox(
            "Availability", ["", "InStock", "OutOfStock", "PreOrder"], key="avail"
        )

ready = bool(title_input and description_input and url_input)
if not ready:
    st.warning("Fill in Title, Description, and URL to enable optimisation.")

optimise_clicked = st.button("Optimise", type="primary", disabled=not ready)

if optimise_clicked:
    working_doc = Document(io.BytesIO(fresh_bytes))

    for idx, level, text in extract_headings(working_doc):
        wrap_heading_text(working_doc.paragraphs[idx], level)

    insert_meta_table_at_top(working_doc, title_input, description_input, url_input)

    if schema_type == "Article":
        primary_schema = build_article_schema(
            title_input, description_input, url_input,
            publisher_name_input, publisher_logo_input, date_published_input,
            image_url=image_input,
        )
    else:
        primary_schema = build_product_schema(
            title_input, description_input, url_input, image_url=image_input,
            brand=brand_input, price=price_input, currency=currency_input,
            availability=availability_input,
        )
    breadcrumb_schema = build_breadcrumb_schema(url_input)
    append_schema_section(working_doc, schema_type, primary_schema, breadcrumb_schema, date_published_input)

    output_stream = io.BytesIO()
    working_doc.save(output_stream)
    output_stream.seek(0)

    st.session_state["optimized_bytes"] = output_stream.getvalue()
    st.session_state["primary_schema_preview"] = primary_schema
    st.session_state["breadcrumb_schema_preview"] = breadcrumb_schema
    st.session_state["schema_type_preview"] = schema_type
    st.session_state["headings_preview"] = extract_headings(source_doc_preview)

if "optimized_bytes" in st.session_state:
    st.divider()
    st.subheader("Preview of changes")

    st.markdown("**Meta table**")
    st.table({
        "Field": ["Meta Title", "Meta Description", "URL"],
        "Value": [title_input, description_input, url_input],
    })

    st.markdown("**Tagged headings**")
    if st.session_state["headings_preview"]:
        for _, level, text in st.session_state["headings_preview"]:
            st.code(f"<h{level}>{text}</h{level}>", language="html")
    else:
        st.write("No headings were found to tag.")

    st.markdown(f"**{st.session_state['schema_type_preview']} Schema (JSON-LD)**")
    st.code(json.dumps(st.session_state["primary_schema_preview"], indent=2), language="json")

    st.markdown("**Breadcrumb Schema (JSON-LD)**")
    st.code(json.dumps(st.session_state["breadcrumb_schema_preview"], indent=2), language="json")

    file_stub = title_input.strip().lower().replace(" ", "-")[:40] or "optimized-doc"
    timestamp = datetime.now().strftime("%Y%m%d")
    st.download_button(
        "Download optimized .docx",
        data=st.session_state["optimized_bytes"],
        file_name=f"{file_stub}-optimized-{timestamp}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    # ----------------------------------------------------------------
    # Anchor Text & URL Checker
    # ----------------------------------------------------------------
    st.divider()
    st.subheader("Anchor Text & URL Checker")
    st.caption("Pulls every hyperlink out of the uploaded fresh doc, flags weak anchor text, and checks if the URL is live.")

    if st.button("Check Anchor Text & URLs"):
        links = extract_hyperlinks(source_doc_preview)
        if not links:
            st.session_state["link_check_rows"] = []
        else:
            with st.spinner(f"Checking {len(links)} link(s)..."):
                with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
                    statuses = list(executor.map(lambda l: check_url_status(l["url"]), links))
            rows = []
            for link, code in zip(links, statuses):
                rows.append({
                    "Anchor Text": link["anchor_text"].strip() or "(empty)",
                    "URL": link["url"],
                    "Status": status_label(code),
                    "Notes": classify_anchor(link["anchor_text"], link["url"]),
                })
            st.session_state["link_check_rows"] = rows

    if "link_check_rows" in st.session_state:
        rows = st.session_state["link_check_rows"]
        if not rows:
            st.write("No hyperlinks found in this document.")
        else:
            broken = sum(1 for r in rows if r["Status"].startswith(("Broken", "Unreachable")))
            weak = sum(1 for r in rows if r["Notes"] != "OK")
            st.write(f"{len(rows)} link(s) found — {broken} broken/unreachable, {weak} with anchor text issues.")
            st.dataframe(rows, use_container_width=True, hide_index=True)